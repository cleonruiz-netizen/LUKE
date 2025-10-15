# core/document_processor.py
import io
import requests
from typing import List, Tuple, Union
from fastapi import UploadFile
import fitz  # PyMuPDF
import docx
from PIL import Image
import pytesseract

pytesseract.pytesseract.tesseract_cmd = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"

class DocumentProcessor:
    """Handles file parsing, text extraction (including OCR), and quality checks."""

    def __init__(self, file: Union[UploadFile, bytes, str]):
        """
        file can be:
        - UploadFile (FastAPI upload)
        - bytes (in-memory file)
        - str (URL or local path)
        """
        self.file_content = None
        self.filename = None

        if isinstance(file, UploadFile):
            self.filename = file.filename
            file.file.seek(0)
            self.file_content = file.file.read()

        elif isinstance(file, bytes):
            self.filename = "unknown.pdf"
            self.file_content = file

        elif isinstance(file, str):
            if file.startswith("http"):  # URL from Supabase
                self.filename = file.split("/")[-1]
                resp = requests.get(file)
                resp.raise_for_status()
                self.file_content = resp.content
            else:  # local path
                self.filename = file.split("/")[-1]
                with open(file, "rb") as f:
                    self.file_content = f.read()

        else:
            raise ValueError("Unsupported file input type")

        if not (self.filename.endswith('.pdf') or self.filename.endswith('.docx')):
            raise ValueError("Unsupported file type. Please provide a PDF or DOCX.")
        

    def process(self) -> Tuple[str, List[str]]:
        if self.filename.endswith('.pdf'):
            full_text, num_pages = self._extract_text_from_pdf()
        elif self.filename.endswith('.docx'):
            full_text, num_pages = self._extract_text_from_docx()
        else:
            full_text, num_pages = "", 0

        quality_flags = self._check_scan_quality(full_text, num_pages)
        return full_text, quality_flags
    

    def _extract_text_from_pdf(self) -> Tuple[str, int]:
        """Extracts all text from a PDF. Falls back to OCR if no text found."""
        text = ""
        doc = fitz.open(stream=self.file_content, filetype="pdf")
        for page in doc:
            page_text = page.get_text().strip()
            if page_text:
                text += page_text + "\n\n"
            else:
                # OCR fallback for scanned pages
                pix = page.get_pixmap()
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                ocr_text = pytesseract.image_to_string(img)
                text += ocr_text + "\n\n"
        return text, doc.page_count

    def _extract_text_from_docx(self) -> Tuple[str, int]:
        """Extracts all text from a DOCX document, including tables."""
        doc = docx.Document(io.BytesIO(self.file_content))
        text = "\n\n".join([para.text for para in doc.paragraphs])

        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"

        num_pages = max(1, len(text) // 3000)  # rough estimate
        return text, num_pages

    def _check_scan_quality(self, text: str, num_pages: int) -> List[str]:
        """
        Implements a guardrail for low OCR confidence.
        If characters per page are very low, flag it as a potential bad scan.
        """
        flags = []
        if num_pages > 0:
            chars_per_page = len(text) / num_pages
            if chars_per_page < 150:  # threshold for likely bad scan
                flags.append("low_ocr_confidence")
        return flags
